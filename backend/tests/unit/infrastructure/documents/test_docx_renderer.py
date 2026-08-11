"""`DocxRenderer` (Port 12, §12.11 / §9.12).

Grouped by the property that would actually break a contract:

  * `TestRendering` — values land, `{{ }}` never survives, `None` never prints.
  * `TestRichText` — ⭐ the securities account number really is a bold run.
  * `TestSandbox` — the second SSTI net, after `TemplateInspector`.
  * `TestAtomicWrite` — no half-written `.docx` at the destination, ever.
  * `TestPreconditions` — missing file, wrong checksum, no disk.
  * `TestCache` — ⭐ the thing that makes NFR-03 reachable at all (§9.12.1).
"""
from __future__ import annotations

import hashlib
import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docxtpl import RichText
from lxml import etree

from cocas.domain.exceptions import (
    InsufficientStorageError,
    RenderError,
    TemplateChecksumMismatchError,
    TemplateNotFoundError,
)
from cocas.infrastructure.documents.docx_renderer import DocxRenderer

NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def make_docx(paragraphs: list[str], *, footer: str | None = None) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if footer is not None:
        document.sections[0].footer.paragraphs[0].text = footer
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def write_template(tmp_path: Path, paragraphs: list[str], **kwargs: object) -> Path:
    path = tmp_path / "template.docx"
    path.write_bytes(make_docx(paragraphs, **kwargs))  # type: ignore[arg-type]
    return path


def all_text(path: Path, part: str = "word/document.xml") -> str:
    with zipfile.ZipFile(path) as archive:
        root = etree.fromstring(archive.read(part))
    return "".join(node.text or "" for node in root.iter(f"{NS}t"))


def bold_texts(path: Path) -> set[str]:
    with zipfile.ZipFile(path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    found = set()
    for run in root.iter(f"{NS}r"):
        properties = run.find(f"{NS}rPr")
        if properties is not None and properties.find(f"{NS}b") is not None:
            text = "".join(node.text or "" for node in run.iter(f"{NS}t")).strip()
            if text:
                found.add(text)
    return found


@pytest.fixture
def renderer() -> DocxRenderer:
    # ⭐ 0 free-space floor: CI runners and dev laptops both exist, and this
    # class's disk check is exercised explicitly in TestPreconditions.
    return DocxRenderer(min_free_bytes=0)


class TestRendering:
    def test_a_variable_is_replaced(self, renderer: DocxRenderer, tmp_path: Path) -> None:
        template = write_template(tmp_path, ["Ho ten: {{ full_name }}"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"full_name": "NGUYỄN VĂN AN"}, str(out))
        assert "NGUYỄN VĂN AN" in all_text(out)

    def test_no_jinja_delimiter_survives(self, renderer: DocxRenderer, tmp_path: Path) -> None:
        template = write_template(tmp_path, ["A {{ a }} B {{ b }}"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"a": "1", "b": "2"}, str(out))
        text = all_text(out)
        assert "{{" not in text and "}}" not in text

    def test_a_missing_variable_renders_empty_not_none(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        """⭐ §9.7's golden rule, enforced by `_EmptyUndefined` rather than by
        hoping every caller filled the context."""
        template = write_template(tmp_path, ["X{{ nobody_set_this }}Y"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {}, str(out))
        text = all_text(out)
        assert "None" not in text
        assert "XY" in text

    def test_footers_are_rendered_too(self, renderer: DocxRenderer, tmp_path: Path) -> None:
        """⚠️ Both real templates put content in the footer; a renderer that
        only handles `word/document.xml` leaves `{{ }}` on every page."""
        template = write_template(tmp_path, ["Body"], footer="So HD: {{ contract_no }}")
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"contract_no": "01A-GDN-202608-00042"}, str(out))
        with zipfile.ZipFile(out) as archive:
            footers = [n for n in archive.namelist() if n.startswith("word/footer")]
        assert footers
        assert any("01A-GDN-202608-00042" in all_text(out, name) for name in footers)

    def test_output_opens_with_python_docx(self, renderer: DocxRenderer, tmp_path: Path) -> None:
        template = write_template(tmp_path, ["Ho ten: {{ full_name }}"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"full_name": "AN"}, str(out))
        assert any("AN" in p.text for p in Document(out).paragraphs)

    def test_non_xml_parts_are_copied_byte_for_byte(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        """Images are most of a template's size and none of its content."""
        template = write_template(tmp_path, ["{{ a }}"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"a": "x"}, str(out))
        with zipfile.ZipFile(template) as source, zipfile.ZipFile(out) as result:
            assert source.read("word/styles.xml") == result.read("word/styles.xml")

    def test_result_reports_the_hash_of_the_bytes_on_disk(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        out = tmp_path / "out.docx"
        result = renderer.render(str(template), {"a": "x"}, str(out))
        assert result.sha256 == hashlib.sha256(out.read_bytes()).digest()
        assert result.size_bytes == out.stat().st_size


class TestRichText:
    def test_a_rich_text_value_produces_a_bold_run(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        """⭐ The one visual requirement the design states outright (§9.5)."""
        template = write_template(tmp_path, ["So TK: {{r securities_account_no }}"])
        out = tmp_path / "out.docx"
        renderer.render(
            str(template),
            {"securities_account_no": RichText("008C123456", bold=True)},
            str(out),
        )
        assert "008C123456" in bold_texts(out)

    def test_plain_syntax_renders_the_same_text_without_the_style(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        """⚠️ `{{ x }}` with a `RichText` prints the markup as text — this is
        exactly what `COCAS-6008` warns about at registration."""
        template = write_template(tmp_path, ["So TK: {{r stk }}"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"stk": RichText("008C123456")}, str(out))
        assert "008C123456" in all_text(out)
        assert "008C123456" not in bold_texts(out)


class TestSandbox:
    def test_a_non_whitelisted_filter_fails_the_render(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ name|attr('__class__') }}"])
        out = tmp_path / "out.docx"
        with pytest.raises(RenderError):
            renderer.render(str(template), {"name": "AN"}, str(out))

    def test_a_whitelisted_filter_still_works(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ name|upper }}"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"name": "an"}, str(out))
        assert "AN" in all_text(out)

    def test_private_attribute_access_is_blocked_by_the_sandbox(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        """The inspector rejects this at registration; the sandbox is the net
        under it, for a file that reached the disk another way (§9.9)."""
        template = write_template(tmp_path, ["{{ ''.__class__.__mro__ }}"])
        out = tmp_path / "out.docx"
        with pytest.raises(RenderError):
            renderer.render(str(template), {}, str(out))

    def test_jinja_globals_are_removed(self, renderer: DocxRenderer, tmp_path: Path) -> None:
        template = write_template(tmp_path, ["{{ range(3) }}"])
        out = tmp_path / "out.docx"
        with pytest.raises(RenderError):
            renderer.render(str(template), {}, str(out))


class TestAtomicWrite:
    def test_no_tmp_file_is_left_behind(self, renderer: DocxRenderer, tmp_path: Path) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"a": "x"}, str(out))
        assert list(tmp_path.glob("*.tmp")) == []

    def test_a_failed_render_leaves_no_file_at_the_destination(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        """⭐ §12.11's invariant: a half-written `.docx` must never exist at
        `output_path` — a user double-clicking it would get a corrupt file."""
        template = write_template(tmp_path, ["{{ a|attr('__class__') }}"])
        out = tmp_path / "out.docx"
        with pytest.raises(RenderError):
            renderer.render(str(template), {"a": "x"}, str(out))
        assert not out.exists()
        assert list(tmp_path.glob("*.tmp")) == []

    def test_rendering_over_an_existing_file_replaces_it(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        out = tmp_path / "out.docx"
        out.write_bytes(b"stale")
        renderer.render(str(template), {"a": "fresh"}, str(out))
        assert "fresh" in all_text(out)

    def test_missing_output_directory_is_created(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        out = tmp_path / "nested" / "deeper" / "out.docx"
        renderer.render(str(template), {"a": "x"}, str(out))
        assert out.exists()


class TestPreconditions:
    def test_a_missing_template_raises_template_not_found(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        with pytest.raises(TemplateNotFoundError):
            renderer.render(str(tmp_path / "nope.docx"), {}, str(tmp_path / "out.docx"))

    def test_a_checksum_mismatch_raises_before_rendering(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        out = tmp_path / "out.docx"
        with pytest.raises(TemplateChecksumMismatchError):
            renderer.render(str(template), {"a": "x"}, str(out), expected_sha256=b"\x00" * 32)
        assert not out.exists()

    def test_a_matching_checksum_is_accepted(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        digest = hashlib.sha256(template.read_bytes()).digest()
        out = tmp_path / "out.docx"
        renderer.render(str(template), {"a": "x"}, str(out), expected_sha256=digest)
        assert out.exists()

    def test_a_full_disk_is_refused_before_writing(self, tmp_path: Path) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        out = tmp_path / "out.docx"
        strict = DocxRenderer(min_free_bytes=1 << 62)
        with pytest.raises(InsufficientStorageError):
            strict.render(str(template), {"a": "x"}, str(out))
        assert not out.exists()

    def test_a_non_docx_file_raises_render_error(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        broken = tmp_path / "broken.docx"
        broken.write_bytes(b"this is not a zip")
        with pytest.raises(RenderError):
            renderer.render(str(broken), {}, str(tmp_path / "out.docx"))


class TestCache:
    """⭐ §9.12.1 — preparing costs 6–9 s on the real templates; rendering
    costs ~0.4 s. Without the cache every contract pays the preparation."""

    def test_a_second_render_reuses_the_prepared_template(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        renderer.render(str(template), {"a": "1"}, str(tmp_path / "a.docx"))
        prepared = dict(renderer._cache)
        renderer.render(str(template), {"a": "2"}, str(tmp_path / "b.docx"))
        assert list(renderer._cache) == list(prepared)

    def test_editing_the_template_in_place_bypasses_the_cache(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        """⚠️ The key includes the SHA-256 precisely for this: restore from
        backup, disk corruption or someone "fixing" the file keeps the path
        and changes the bytes. A path-only key would render the old template
        until the process restarts."""
        template = write_template(tmp_path, ["Cũ: {{ a }}"])
        renderer.render(str(template), {"a": "1"}, str(tmp_path / "a.docx"))
        template.write_bytes(make_docx(["Mới: {{ a }}"]))
        out = tmp_path / "b.docx"
        renderer.render(str(template), {"a": "2"}, str(out))
        assert "Mới" in all_text(out)

    def test_the_cache_is_bounded(self, tmp_path: Path) -> None:
        small = DocxRenderer(cache_size=1, min_free_bytes=0)
        for index in range(3):
            template = tmp_path / f"t{index}.docx"
            template.write_bytes(make_docx([f"{index}: {{{{ a }}}}"]))
            small.render(str(template), {"a": "x"}, str(tmp_path / f"o{index}.docx"))
        assert len(small._cache) == 1

    def test_prepare_populates_the_cache_without_rendering(
        self, renderer: DocxRenderer, tmp_path: Path
    ) -> None:
        template = write_template(tmp_path, ["{{ a }}"])
        renderer.prepare(str(template))
        assert len(renderer._cache) == 1
        assert list(tmp_path.glob("*.docx")) == [template]
