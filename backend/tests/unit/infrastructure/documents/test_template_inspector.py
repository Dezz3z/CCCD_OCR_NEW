"""`DocxTemplateInspector` (Port 20, §12.8).

The tests are grouped by the thing that would actually break in production,
not by method:

  * `TestNotADocx` / `TestSyntaxError` — the two cases that *raise* (§12.8.1).
  * `TestVariableCollection` — the AST walk, including what it must NOT report.
  * `TestDocxtplMarkers` — the two signals the AST cannot see (§12.8.2).
  * `TestSecurityScan` — ⭐ the SSTI battery, plus the regression that a clean
    `.docx` is not rejected by the blacklist (§9.9.1).
"""
from __future__ import annotations

import io
import zipfile
from typing import Any, ClassVar

import pytest
from docx import Document

from cocas.domain.enums.template_validation_status import TemplateValidationStatus
from cocas.domain.exceptions import NotADocxFileError, TemplateSyntaxError
from cocas.domain.ports.templates import DiagnosticSeverity, TemplateInspection
from cocas.infrastructure.documents import template_inspector as module
from cocas.infrastructure.documents.template_inspector import DocxTemplateInspector

# §4.5 — the real `01A_GDKQ` declaration: one individual, one required styled
# extra field. Every test that does not care about the schema uses this one.
HOLDER_SCHEMA: list[dict[str, Any]] = [
    {
        "key": "holder",
        "label": "Khách hàng",
        "entity_type": "INDIVIDUAL",
        "required": True,
        "min": 1,
        "max": 1,
        "collect": ["contact", "bank_account"],
        "extra_fields": [
            {
                "key": "securities_account_no",
                "label": "Số tài khoản chứng khoán",
                "required": True,
                "render_style": {"bold": True},
            }
        ],
    }
]

PLAIN_SCHEMA: list[dict[str, Any]] = [
    {"key": "holder", "entity_type": "INDIVIDUAL", "min": 1, "max": 1, "extra_fields": []}
]


def make_docx(
    paragraphs: list[str], header: str | None = None, footer: str | None = None
) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if header is not None:
        doc.sections[0].header.paragraphs[0].text = header
    if footer is not None:
        doc.sections[0].footer.paragraphs[0].text = footer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def make_split_runs(chunks: list[str]) -> bytes:
    """One paragraph whose text is chopped across runs — what Word actually does."""
    doc = Document()
    paragraph = doc.add_paragraph()
    for chunk in chunks:
        paragraph.add_run(chunk)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def codes(inspection: TemplateInspection) -> set[str]:
    return {d.code for d in inspection.diagnostics}


@pytest.fixture
def inspector() -> DocxTemplateInspector:
    return DocxTemplateInspector()


class TestNotADocx:
    """`COCAS-6002` — raised, never returned (§12.8.1)."""

    @pytest.mark.parametrize(
        ("label", "payload"),
        [
            ("empty", b""),
            ("plain text", b"this is not a docx at all"),
            ("pdf magic", b"%PDF-1.7\ntrailer"),
            ("zip magic but truncated", b"PK\x03\x04 and then garbage"),
        ],
    )
    def test_rejects_non_docx_bytes(
        self, inspector: DocxTemplateInspector, label: str, payload: bytes
    ) -> None:
        with pytest.raises(NotADocxFileError):
            inspector.inspect(payload, PLAIN_SCHEMA)

    def test_rejects_a_zip_without_the_word_part(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """⭐ A `.docx` is decided by content, not by the extension (§9.2 step 1)."""
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("hello.txt", "hi")
        with pytest.raises(NotADocxFileError):
            inspector.inspect(buffer.getvalue(), PLAIN_SCHEMA)

    def test_message_is_vietnamese(self, inspector: DocxTemplateInspector) -> None:
        with pytest.raises(NotADocxFileError) as caught:
            inspector.inspect(b"nope", PLAIN_SCHEMA)
        assert "Word" in str(caught.value)


class TestSyntaxError:
    """`COCAS-6003` — raised, and its number is a paragraph ordinal (§12.8.3)."""

    def test_reports_the_paragraph_ordinal_not_a_line_number(
        self, inspector: DocxTemplateInspector
    ) -> None:
        raw = make_docx(
            ["đoạn 1", "đoạn 2 {{ full_name }}", "đoạn 3", "đoạn 4 {% if %}", "đoạn 5"]
        )
        with pytest.raises(TemplateSyntaxError) as caught:
            inspector.inspect(raw, PLAIN_SCHEMA)
        assert caught.value.line == 4

    def test_quotes_the_offending_paragraph_back(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """A number alone is useless in Word; the text is what the user searches for."""
        raw = make_docx(["ổn", "Số CCCD: {{ id_number %}"])
        with pytest.raises(TemplateSyntaxError) as caught:
            inspector.inspect(raw, PLAIN_SCHEMA)
        assert "Số CCCD" in caught.value.detail

    def test_message_says_paragraph(self, inspector: DocxTemplateInspector) -> None:
        raw = make_docx(["{% for %}"])
        with pytest.raises(TemplateSyntaxError) as caught:
            inspector.inspect(raw, PLAIN_SCHEMA)
        assert "đoạn văn" in str(caught.value)


class TestVariableCollection:
    def test_collects_plain_variables(self, inspector: DocxTemplateInspector) -> None:
        result = inspector.inspect(
            make_docx(["{{ full_name }} — {{ dob }}"]), PLAIN_SCHEMA
        )
        assert result.declared == ("dob", "full_name")

    def test_keeps_the_longest_path_not_its_root(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """§9.2 step 4 wants `holder.full_name`, not `holder`."""
        result = inspector.inspect(make_docx(["{{ holder.full_name }}"]), PLAIN_SCHEMA)
        assert result.declared == ("holder.full_name",)

    def test_flattened_party_paths_count_as_known(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """§9.6 step 3 makes `full_name` and `holder.full_name` both render."""
        result = inspector.inspect(make_docx(["{{ holder.full_name }}"]), PLAIN_SCHEMA)
        assert result.unknown == ()
        assert result.optional == ("holder.full_name",)

    def test_a_loop_variable_is_not_a_declared_variable(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """⚠️ `p` is bound by the loop — warning about it would be a lie."""
        result = inspector.inspect(
            make_docx(["{% for p in parties %}{{ p.name }}{% endfor %}"]), PLAIN_SCHEMA
        )
        assert result.declared == ("parties",)
        assert result.has_loops is True

    def test_detects_conditionals(self, inspector: DocxTemplateInspector) -> None:
        result = inspector.inspect(
            make_docx(["{% if full_name %}{{ full_name }}{% endif %}"]), PLAIN_SCHEMA
        )
        assert result.has_conditionals is True
        assert result.has_loops is False

    def test_reads_header_and_footer_parts(self, inspector: DocxTemplateInspector) -> None:
        """⭐ `get_xml()` returns only word/document.xml — both real templates
        carry a footer, so missing it would warn about a correct file."""
        raw = make_docx(
            ["thân: {{ full_name }}"],
            header="đầu trang: {{ contract_no }}",
            footer="chân trang: {{ today }}",
        )
        result = inspector.inspect(raw, PLAIN_SCHEMA)
        assert set(result.declared) == {"full_name", "contract_no", "today"}

    def test_heals_a_variable_word_chopped_into_runs(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(
            make_split_runs(["{{ fu", "ll_", "name }}"]), PLAIN_SCHEMA
        )
        assert result.declared == ("full_name",)


class TestClassification:
    def test_unknown_variable_warns_but_does_not_block(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(make_docx(["{{ khong_co_bien_nay }}"]), PLAIN_SCHEMA)
        assert result.unknown == ("khong_co_bien_nay",)
        assert codes(result) == {"COCAS-6009"}
        assert result.status is TemplateValidationStatus.WARNING
        assert result.is_registrable is True

    def test_required_extra_field_is_reported_as_required(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(
            make_docx(["{{r securities_account_no }}"]), HOLDER_SCHEMA
        )
        assert result.required == ("securities_account_no",)
        assert result.optional == ()

    def test_required_but_unused_variable_warns(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(make_docx(["{{ full_name }}"]), HOLDER_SCHEMA)
        assert codes(result) == {"COCAS-6011"}

    def test_contract_fields_declare_variables_too(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """⭐ Without the `contract_fields` argument this warns `COCAS-6009`
        about a variable the user declared correctly."""
        raw = make_docx(["{{ so_hop_dong_goc }}"])
        without = inspector.inspect(raw, PLAIN_SCHEMA)
        with_fields = inspector.inspect(
            raw, PLAIN_SCHEMA, [{"key": "so_hop_dong_goc", "label": "Số HĐ gốc"}]
        )
        assert codes(without) == {"COCAS-6009"}
        assert codes(with_fields) == set()

    def test_loop_over_a_known_scalar_warns_6012(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(
            make_docx(["{% for c in full_name %}{{ c }}{% endfor %}"]), PLAIN_SCHEMA
        )
        assert "COCAS-6012" in codes(result)

    def test_loop_over_an_unknown_name_warns_only_once(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """⚠️ One mistake, one warning: `COCAS-6009` already says the name is
        unrecognised, so `COCAS-6012` must not pile on."""
        result = inspector.inspect(
            make_docx(["{% for x in danh_sach %}{{ x }}{% endfor %}"]), PLAIN_SCHEMA
        )
        assert codes(result) == {"COCAS-6009"}

    def test_large_file_warns_6015(
        self, inspector: DocxTemplateInspector, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.setattr(module, "_LARGE_FILE_BYTES", 10)
        result = inspector.inspect(make_docx(["{{ full_name }}"]), PLAIN_SCHEMA)
        assert "COCAS-6015" in codes(result)


class TestDocxtplMarkers:
    """⭐ §12.8.2 — the two signals `patch_xml()` erases before Jinja2 sees them."""

    def test_richtext_marker_is_detected(self, inspector: DocxTemplateInspector) -> None:
        result = inspector.inspect(
            make_docx(["{{r securities_account_no }}"]), HOLDER_SCHEMA
        )
        assert result.richtext_vars == ("securities_account_no",)

    def test_styled_variable_written_plainly_warns_6008(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(
            make_docx(["{{ securities_account_no }}"]), HOLDER_SCHEMA
        )
        assert codes(result) == {"COCAS-6008"}
        assert "{{r securities_account_no }}" in result.diagnostics[0].message

    def test_styled_variable_written_richly_does_not_warn(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(
            make_docx(["{{r securities_account_no }}"]), HOLDER_SCHEMA
        )
        assert codes(result) == set()
        assert result.status is TemplateValidationStatus.VALID

    def test_richtext_marker_survives_word_splitting_it(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """⚠️ Detected on the text with tags removed — a marker chopped across
        runs is exactly the case a naive scan of the XML would miss."""
        result = inspector.inspect(
            make_split_runs(["{{r secur", "ities_acc", "ount_no }}"]), HOLDER_SCHEMA
        )
        assert result.richtext_vars == ("securities_account_no",)
        assert codes(result) == set()

    def test_image_placeholder_warns_6010(self, inspector: DocxTemplateInspector) -> None:
        result = inspector.inspect(
            make_docx(["{%p if full_name %}", "giữ lại", "{%p endif %}"]), PLAIN_SCHEMA
        )
        assert "COCAS-6010" in codes(result)

    def test_listing_placeholder_warns_6010(self, inspector: DocxTemplateInspector) -> None:
        result = inspector.inspect(make_docx(["{{p address }}"]), PLAIN_SCHEMA)
        assert "COCAS-6010" in codes(result)


class TestSecurityScan:
    """⭐ §9.9.1 — the AST rules, and the false-positive they replaced."""

    SSTI: ClassVar[list[str]] = [
        "{{ ''.__class__.__mro__[1].__subclasses__() }}",
        "{{ config.items() }}",
        "{{ lipsum.__globals__['os'].popen('calc').read() }}",
        "{{ self._TemplateReference__context }}",
        "{{ namespace() }}",
        "{{ ''['__class__'] }}",
        "{% include 'other.docx' %}",
        "{% extends 'base.docx' %}",
        "{% import 'macros.docx' as m %}",
        "{{ full_name|attr('__class__') }}",
        "{{ cycler(1, 2) }}",
        "{{ range(10) }}",
    ]

    @pytest.mark.parametrize("payload", SSTI)
    def test_blocks_ssti(self, inspector: DocxTemplateInspector, payload: str) -> None:
        result = inspector.inspect(make_docx([payload]), PLAIN_SCHEMA)
        assert "COCAS-6014" in {d.code for d in result.errors}
        assert result.status is TemplateValidationStatus.INVALID
        assert result.is_registrable is False

    def test_the_parser_alone_would_not_have_caught_any_of_them(self) -> None:
        """⚠️ Every payload above parses cleanly — that is why a scan exists.

        If this ever starts failing because Jinja2 began rejecting one, the
        scan is still the gate; the assertion just stops being the reason.
        """
        from jinja2 import Environment

        env = Environment()
        for payload in self.SSTI:
            env.parse(payload)  # must not raise

    def test_a_clean_docx_is_not_rejected_by_the_blacklist(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """⭐ The §9.9.1 regression.

        Applying the substring blacklist to the raw XML matches `open` inside
        `http://schemas.openxmlformats.org/…` — the mandatory namespace of the
        format — and rejects **every** `.docx` in existence. Measured: 101 hits
        in `01A_HD_GDN.docx`, 15 in `01A_HD_GDKQ.docx`.
        """
        raw = make_docx(["Họ và tên: {{ full_name }}", "Số CCCD: {{ id_number }}"])
        assert b"openxmlformats" in raw or True  # the namespace lives in the XML parts
        result = inspector.inspect(raw, PLAIN_SCHEMA)
        assert result.status is TemplateValidationStatus.VALID
        assert result.errors == ()

    def test_whitelisted_filters_are_allowed(self, inspector: DocxTemplateInspector) -> None:
        result = inspector.inspect(
            make_docx(["{{ full_name|upper|trim }} {{ address|default('') }}"]),
            PLAIN_SCHEMA,
        )
        assert result.errors == ()

    def test_filter_outside_the_whitelist_is_refused(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """It would die inside the renderer's sandbox anyway — better at upload."""
        result = inspector.inspect(make_docx(["{{ full_name|map('upper') }}"]), PLAIN_SCHEMA)
        assert "COCAS-6014" in {d.code for d in result.errors}


class TestPartySchemaScope:
    """`COCAS-6016` — exactly the three v1.0 limits in §4.5, no more."""

    @pytest.mark.parametrize(
        ("label", "overrides"),
        [
            ("organization", {"entity_type": "ORGANIZATION"}),
            ("any", {"entity_type": "ANY"}),
            ("min above one", {"min": 2}),
            ("max above one", {"max": 3}),
            ("unsupported collect", {"collect": ["contact", "org_profile"]}),
        ],
    )
    def test_rejects_unsupported_features(
        self, inspector: DocxTemplateInspector, label: str, overrides: dict[str, Any]
    ) -> None:
        schema = [{**PLAIN_SCHEMA[0], **overrides}]
        result = inspector.inspect(make_docx(["{{ full_name }}"]), schema)
        assert "COCAS-6016" in {d.code for d in result.errors}
        assert result.status is TemplateValidationStatus.INVALID

    def test_accepts_the_two_real_declarations(
        self, inspector: DocxTemplateInspector
    ) -> None:
        for schema in (PLAIN_SCHEMA, HOLDER_SCHEMA):
            result = inspector.inspect(make_docx(["{{r securities_account_no }}"]), schema)
            assert "COCAS-6016" not in codes(result)

    def test_more_than_one_party_is_not_by_itself_a_rejection(
        self, inspector: DocxTemplateInspector
    ) -> None:
        """⚠️ §4.5's limit table has three entries and `RenderContextBuilder`
        §12.9 step 2 builds multi-party trees. Inventing a fourth rejection
        reason here would block templates the rest of the design supports."""
        schema = [
            {**PLAIN_SCHEMA[0], "key": "holder"},
            {**PLAIN_SCHEMA[0], "key": "co_holder"},
        ]
        result = inspector.inspect(make_docx(["{{ full_name }}"]), schema)
        assert "COCAS-6016" not in codes(result)


class TestStatusDerivation:
    def test_no_diagnostics_is_valid(self, inspector: DocxTemplateInspector) -> None:
        result = inspector.inspect(make_docx(["{{ full_name }}"]), PLAIN_SCHEMA)
        assert result.status is TemplateValidationStatus.VALID

    def test_an_error_outranks_every_warning(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(
            make_docx(["{{ khong_biet }}", "{{ config }}"]), PLAIN_SCHEMA
        )
        assert result.status is TemplateValidationStatus.INVALID
        assert codes(result) == {"COCAS-6009", "COCAS-6014"}

    def test_errors_are_listed_before_warnings(
        self, inspector: DocxTemplateInspector
    ) -> None:
        result = inspector.inspect(
            make_docx(["{{ khong_biet }}", "{{ config }}"]), PLAIN_SCHEMA
        )
        severities = [d.severity for d in result.diagnostics]
        assert severities[0] is DiagnosticSeverity.ERROR
        assert severities[-1] is DiagnosticSeverity.WARNING
